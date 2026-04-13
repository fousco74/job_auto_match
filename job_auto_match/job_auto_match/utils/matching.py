import frappe
import requests
from google import genai
from google.genai import types
from frappe.utils.file_manager import get_file_path
from urllib.parse import urljoin
import time
import random
from google.genai import errors as genai_errors
from jinja2 import TemplateNotFound
import json, mimetypes, pathlib, tempfile, shutil, subprocess
import mammoth
from bs4 import BeautifulSoup
from docx import Document
from docx.opc.exceptions import PackageNotFoundError


_SETTINGS_DOCTYPE = "Job Matching Integration Settings"



# --- Flags mapping (adapte si tes fieldnames diffèrent) ---
FLAG_MATCHING_IN_PROGRESS = "custom_is_matching_in_progress"
FLAG_MATCHING_FAILED      = "custom_is_matching_failed"
FLAG_NOT_MATCH_EMAIL_SENT = "custom_not_match_email_sent"
FLAG_INVITES_SENT         = "custom_invites_sent"
FIELD_AI_LAST_ERROR       = "custom_ai_last_error"
FIELD_STATUT              = "custom_status"


# ---------- Helpers: Word→PDF/texte & préparation des "parts" ----------

def _libreoffice_to_pdf_bytes(input_path: str) -> bytes | None:
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice:
        return None
    outdir = tempfile.mkdtemp(prefix="lo2pdf_")
    try:
        cmd = [soffice, "--headless", "--convert-to", "pdf", "--outdir", outdir, input_path]
        subprocess.run(cmd, check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
        base = pathlib.Path(input_path).stem
        pdf_path = pathlib.Path(outdir) / f"{base}.pdf"
        if not pdf_path.exists():
            pdfs = list(pathlib.Path(outdir).glob("*.pdf"))
            if not pdfs:
                return None
            pdf_path = max(pdfs, key=lambda p: p.stat().st_mtime)
        return pdf_path.read_bytes()
    except Exception:
        return None


def _extract_text_from_docx(input_path: str) -> str | None:
    """
    Extract text from a .docx file.
    Returns a single string or None if nothing could be extracted.
    """
    # 1) python-docx path (fast & accurate for .docx)
    try:
        doc = Document(input_path)
        chunks = []

        # paragraphs
        chunks.extend(p.text.strip() for p in doc.paragraphs if p.text and p.text.strip())

        # tables
        for tbl in doc.tables:
            for row in tbl.rows:
                row_txt = " | ".join((c.text or "").strip() for c in row.cells)
                if row_txt.strip():
                    chunks.append(row_txt)

        text = "\n".join(chunks).strip()
        if text:
            return text
    except PackageNotFoundError:
        pass
    except Exception:
        pass

    # 2) Fallback: mammoth (.docx → HTML → plain text)
    try:
        with open(input_path, "rb") as f:
            html = mammoth.convert_to_html(f).value
        soup = BeautifulSoup(html, "html.parser")
        text = soup.get_text("\n").strip()
        return text or None
    except Exception:
        return None


def _prepare_resume_parts_for_gemini(file_path: str):
    p = pathlib.Path(file_path)
    if not p.exists():
        raise FileNotFoundError(f"Fichier introuvable: {file_path}")

    mime, _ = mimetypes.guess_type(file_path)
    mime = mime or "application/octet-stream"

    allowed = {
        "application/pdf",
        "application/msword",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    }
    if mime not in allowed:
        raise ValueError("Format non supporté. Seuls PDF et Word (DOC/DOCX) sont acceptés.")

    # PDF → send as-is
    if mime == "application/pdf":
        data = p.read_bytes()
        parts = [types.Part.from_bytes(data=data, mime_type="application/pdf")]
        return parts, {"strategy": "pdf-inline", "mime": mime}

    # DOCX
    if mime == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        pdf_bytes = _libreoffice_to_pdf_bytes(file_path)
        if pdf_bytes:
            parts = [types.Part.from_bytes(data=pdf_bytes, mime_type="application/pdf")]
            return parts, {"strategy": "word->pdf(lo)", "mime": mime}

        text = _extract_text_from_docx(file_path)
        if text:
            return [text], {"strategy": "word->text", "mime": mime}

        raise RuntimeError("Impossible d'extraire le texte du DOCX.")

    # DOC (legacy)
    if mime == "application/msword":
        pdf_bytes = _libreoffice_to_pdf_bytes(file_path)
        if pdf_bytes:
            parts = [types.Part.from_bytes(data=pdf_bytes, mime_type="application/pdf")]
            return parts, {"strategy": "doc->pdf(lo)", "mime": mime}
        raise RuntimeError(
            "Impossible de convertir le fichier .doc. Installez LibreOffice (soffice) ou fournissez un PDF/DOCX."
        )


def _set_flag(doc, fieldname: str, value: int):
    try:
        if hasattr(doc, fieldname):
            setattr(doc, fieldname, value if value else 0)
        else:
            frappe.logger().warning(f"[FLAGS] Champ manquant sur Job Applicant: {fieldname}")
    except Exception:
        pass


def _set_text(doc, fieldname: str, value: str, max_len=1000):
    try:
        if hasattr(doc, fieldname):
            setattr(doc, fieldname, (value or "")[:max_len])
    except Exception:
        pass


def _set_statut(doc, statut: str):
    """Met à jour custom_status et workflow_state simultanément."""
    _set_text(doc, FIELD_STATUT, statut)
    try:
        doc.workflow_state = statut
    except Exception:
        pass


def _safe_log_error(title: str, err: Exception):
    safe_title = (title or "")[:140]
    try:
        full = f"{frappe.get_traceback()}\n{repr(err)}"
        frappe.log_error(message=full, title=safe_title)
    except Exception:
        try:
            frappe.logger().error(f"[LOG_ERROR_FAIL] {safe_title} :: {err}")
        except Exception:
            pass



def _save_and_reload(doc):
    """
    Sauvegarde depuis un job background sans dépendre d'un compte utilisateur.

    Pourquoi le monkey-patch :
      doc.save(ignore_permissions=True) bypasse les checks de permission doctype,
      MAIS Frappe appelle quand même self.validate_workflow() dans _validate().
      validate_workflow() appelle check_permission("read") sur doc._doc_before_save
      via frappe.session.user — ce qui peut échouer dans un contexte background job.
      En remplaçant validate_workflow au niveau de l'instance, on court-circuite
      cette chaîne sans modifier Frappe ni dépendre d'Administrator.
    """
    doc.flags.ignore_permissions = True
    doc.validate_workflow = lambda: None  # instance attribute → shadowe la méthode de classe
    doc.save(ignore_permissions=True)
    frappe.db.commit()
    doc.reload()

# -----------------------------------------------------------------------------

SAFE_MODEL_CANDIDATES = ["gemini-2.5-pro", "gemini-2.5-flash", "gemini-2.5-flash-lite"]


def call_gemini_with_retry(
    client,
    parts,
    model_candidates=None,
    max_attempts=5,
    base_sleep=1.0,
    max_sleep=10.0,
):
    """
    Appelle Gemini avec retries exponentiels + fallback de modèles.
    - 404/NOT_FOUND ou 'not supported for generateContent' => skip ce modèle
    - 429/5xx/UNAVAILABLE/overloaded/quota => retry avec backoff
    """
    models = model_candidates or SAFE_MODEL_CANDIDATES
    last_exc = None

    for model in models:
        for attempt in range(1, max_attempts + 1):
            try:
                return client.models.generate_content(model=model, contents=parts)

            except genai_errors.APIError as e:
                code = getattr(e, "status_code", None)
                resp = getattr(e, "response_json", {}) or {}
                err = (resp.get("error") or {})
                status = (err.get("status") or "").upper()
                message = (err.get("message") or str(e))

                not_supported = (
                    code in (400, 404) and (
                        "NOT_FOUND" in status
                        or "not found" in message.lower()
                        or "not supported for generatecontent" in message.lower()
                    )
                )
                if not_supported:
                    frappe.logger().warning(
                        f"[GEMINI] Skip modèle non supporté: '{model}' ({code}/{status}) : {message}"
                    )
                    last_exc = e
                    break

                retryable = (
                    code in (429, 500, 502, 503, 504)
                    or "UNAVAILABLE" in status
                    or "RESOURCE_EXHAUSTED" in status
                    or "overloaded" in message.lower()
                    or "quota" in message.lower()
                )
                if retryable and attempt < max_attempts:
                    sleep = min(base_sleep * (2 ** (attempt - 1)), max_sleep) + random.uniform(0, 0.5)
                    frappe.logger().warning(
                        f"[GEMINI] {model} tentative {attempt}/{max_attempts} échouée "
                        f"({code or status}) : {message}. Retry dans {sleep:.1f}s"
                    )
                    time.sleep(sleep)
                    continue

                last_exc = e
                break

            except Exception as e:
                last_exc = e
                break

    raise last_exc


# -----------------------------
# 1) Email Candidat Non Matching
# -----------------------------

def send_candidate_not_matching_email(doc):
    TEMPLATE_PATH = "job_auto_match/templates/emails/candidate_not_matching.html"
    try:
        settings = frappe.get_cached_doc(_SETTINGS_DOCTYPE)

        recipient = (getattr(doc, "email_id", "") or "").strip()
        if not recipient:
            frappe.logger().error("[NOT_MATCH_MAIL] Aucun destinataire renseigné")
            return

        ctx = {
            "applicant_name": getattr(doc, "applicant_name", ""),
            "job_title": getattr(doc, "custom_nom_de_loffre", ""),
            "score": getattr(doc, "custom_matching_score", None),
            "justification": getattr(doc, "custom_justification", "") or "",
        }

        subject_src = getattr(settings, "candidate_not_matching_subject", "") or ""
        subject = frappe.render_template(subject_src, ctx) if subject_src.strip() else frappe._(
            "Votre candidature n'est pas retenue pour le moment"
        )

        header_src = getattr(settings, "candidate_not_matching_email_header", "") or ""
        header = frappe.render_template(header_src, ctx) if header_src.strip() else frappe._(
            "Information sur votre candidature"
        )

        html_src = getattr(settings, "candidate_not_matching_email_template", "") or ""
        if html_src.strip():
            body = frappe.render_template(html_src, ctx)
        else:
            body = frappe.get_template(TEMPLATE_PATH).render(ctx)

        frappe.sendmail(
            recipients=[recipient],
            subject=subject,
            message=body,
            header=header,
        )

        _set_flag(doc, FLAG_NOT_MATCH_EMAIL_SENT, 1)
        _save_and_reload(doc)  # ← reload après save
        frappe.logger().info(f"[NOT_MATCH_MAIL] Email envoyé à {recipient} avec succès.")

    except TemplateNotFound as e:
        _safe_log_error("[NOT_MATCH_MAIL] Template introuvable", e)
        _set_text(doc, FIELD_AI_LAST_ERROR, f"Template introuvable: {e}")
        _set_flag(doc, FLAG_NOT_MATCH_EMAIL_SENT, 0)
        _save_and_reload(doc)  # ← reload après save

    except Exception as e:
        _safe_log_error("[NOT_MATCH_MAIL] Erreur d'envoi", e)
        _set_text(doc, FIELD_AI_LAST_ERROR, f"Email non retenu: {e}")
        _set_flag(doc, FLAG_NOT_MATCH_EMAIL_SENT, 0)
        _save_and_reload(doc)  # ← reload après save


# -----------------------------
# 2) Invitation Testlify
# -----------------------------

def send_candidate_invite(doc, assessments: list) -> list:
    results = []
    any_success = False

    try:
        settings = frappe.get_cached_doc(_SETTINGS_DOCTYPE)

        base_url  = (settings.testlify_base_url or "").strip()
        inv_path  = (settings.testlify_candidate_invite or "").strip()
        token     = (settings.get_password("testlify_token") or "").strip()

        missing = []
        if not base_url: missing.append("URL de base Testlify")
        if not inv_path: missing.append("Endpoint invitation candidat")
        if not token:    missing.append("Token API Testlify")
        if missing:
            raise ValueError(
                f"Configuration Testlify incomplète — champs manquants : {', '.join(missing)}"
            )

        api_url = urljoin(base_url.rstrip('/') + '/', inv_path.lstrip('/'))

        frappe.logger().info(f"[INVITE] URL : {api_url} | Candidat : {doc.custom_first_name} {doc.custom_last_name} <{getattr(doc, 'email_id', None)}>")

        headers = {
            'Authorization': f'Bearer {token}',
            'Content-Type': 'application/json'
        }

        for assessment in (assessments or []):
            if isinstance(assessment, dict):
                assessment_id = assessment.get('id')
                assessment_name = assessment.get('assessment_name')
            else:
                assessment_id = getattr(assessment, 'id', None)
                assessment_name = getattr(assessment, 'assessment_name', None)

            if not assessment_id:
                frappe.logger().warning(f"[INVITE] assessment_id vide, ignoré : {assessment!r}")
                continue

            first_name = (getattr(doc, "custom_first_name", "") or "").strip()
            last_name  = (getattr(doc, "custom_last_name",  "") or "").strip()
            email      = (getattr(doc, "email_id",          "") or "").strip()

            if not email:
                frappe.logger().warning(f"[INVITE] Email candidat vide, assessment {assessment_id} ignoré.")
                results.append({"assessment_id": assessment_id, "status": "skipped", "message": "email vide"})
                continue

            payload = {
                'candidateInvites': [{
                    'firstName': first_name,
                    'lastName':  last_name,
                    'email':     email,
                }],
                'assessmentId': assessment_id,
            }

            response = requests.post(api_url, json=payload, headers=headers, timeout=30)
            status = response.status_code
            try:
                body = response.json()
            except ValueError:
                body = {}

            if status == 200:
                any_success = True
                # Mettre à jour le row existant plutôt qu'en ajouter un doublon
                if hasattr(doc, "custom_assessments"):
                    for row in (doc.custom_assessments or []):
                        if getattr(row, "assessment_id", None) == assessment_id:
                            row.sent = 1
                            break
                    else:
                        doc.append("custom_assessments", {
                            "assessment_id":   assessment_id,
                            "assessment_name": assessment_name,
                            "sent":            1,
                        })
                results.append({"assessment_id": assessment_id, "status": "success"})
            else:
                err_body = body.get("error") if isinstance(body, dict) else None
                msg = (err_body.get("message") if isinstance(err_body, dict) else None) or response.text
                results.append({"assessment_id": assessment_id, "status": status, "message": msg})

        frappe.logger().info(f"[INVITE] Résultats : {results}")
        if any_success:
            _set_flag(doc, FLAG_INVITES_SENT, 1)
        else:
            _set_flag(doc, FLAG_INVITES_SENT, 0)
        _save_and_reload(doc)  # ← reload après save final de la boucle

    except Exception as e:
        _safe_log_error("[INVITE] Erreur fatale", e)
        _set_text(doc, FIELD_AI_LAST_ERROR, f"Invite: {e}")
        _set_flag(doc, FLAG_INVITES_SENT, 0)
        _save_and_reload(doc)  # ← reload après save dans le handler d'erreur
        results.append({"status": "fatal_error", "message": str(e)})

    return results


# -----------------------------
# 3) Process Matching Candidat
# -----------------------------

def process_job_applicant_matching(applicant_name, **kwargs):
    settings = frappe.get_cached_doc(_SETTINGS_DOCTYPE)
    API_KEY = (settings.get_password("gemini_api_key") or "").strip()
    if not API_KEY:
        raise ValueError(
            "Clé API Gemini non configurée dans Job Matching Integration Settings."
        )
    client = genai.Client(api_key=API_KEY)

    qualified_status = settings.status_qualified or "En Cours de qualification"
    status_not_qualified = settings.status_not_qualified or "Top Profil"
    qualification_score_threshold = settings.qualification_score_threshold or 70
    status_rejected = settings.status_rejected or "Rejecté"
    rejected_score = settings.rejected_max_score or 40
    gemini_error_status = settings.gemini_error_status or "Open"

    doc = frappe.get_doc("Job Applicant", applicant_name)

    # ▶️ Flags: démarrage matching
    _set_flag(doc, FLAG_MATCHING_IN_PROGRESS, 1)
    _set_flag(doc, FLAG_MATCHING_FAILED, 0)
    _set_text(doc, FIELD_AI_LAST_ERROR, "")
    _save_and_reload(doc)

    _matching_error = ""       # non-vide = erreur fatale → persistée dans le finally
    _matching_succeeded = False  # True = traitement normal terminé

    try:
        fiche = frappe.get_doc("Job Opening", doc.job_title or "")
        site_url = settings.site_url

        if fiche.custom_active_cv_auto_matching == 0:
            raise ValueError("Job auto match desable")

        if not doc.resume_attachment:
            raise ValueError("Le candidat n'a pas de pièce jointe 'resume_attachment'")

        file_path = get_file_path(doc.resume_attachment)
        if not pathlib.Path(file_path).exists():
            raise FileNotFoundError(f"Fichier introuvable: {file_path}")

        # 🔐 Limiter aux PDF/Word + préparer parts sûrs pour Gemini
        try:
            parts_cv, prep_info = _prepare_resume_parts_for_gemini(file_path)
        except ValueError as bad_fmt:
            _set_flag(doc, FLAG_MATCHING_FAILED, 0)
            _set_flag(doc, FLAG_MATCHING_IN_PROGRESS, 0)
            _set_text(doc, FIELD_AI_LAST_ERROR, str(bad_fmt))
            _set_statut(doc, status_rejected or "Rejecté")
            doc.custom_justification = "Rejeté: format non supporté (PDF ou Word uniquement)."
            doc.custom_matching_score = 0
            doc.applicant_rating = 0.0
            _save_and_reload(doc)
            return

        # --- Étape 0 : vérifier que le document est bien un CV ---
        cv_check_prompt = """
        Tu es un classificateur. Dis si le document fourni est un CV (curriculum vitae/résumé).
        Réponds STRICTEMENT en JSON sans markdown:
        {"is_cv": true|false, "reason": "<raison brève en français>"}
        Critères d'un CV: identité/contact, expériences ou projets professionnels, compétences/outils, éducation/diplômes.
        Exemples NON CV: facture, attestation, lettre simple, offre d'emploi, tract publicitaire, rapport sans section expérience personnelle, photo seule, etc.
        Ne fais aucune supposition si les indices sont absents.
        """
        try:
            cv_check_resp = call_gemini_with_retry(client, parts_cv + [cv_check_prompt])
            try:
                cv_check = json.loads(cv_check_resp.text)
            except Exception:
                cv_check = json.loads(cv_check_resp.text.strip("```json").strip("```").strip())
        except Exception as e:
            cv_check = {"is_cv": True, "reason": "classification sautée (moteur indisponible)"}

        if not cv_check.get("is_cv", True):
            _set_flag(doc, FLAG_MATCHING_FAILED, 0)
            _set_flag(doc, FLAG_MATCHING_IN_PROGRESS, 0)
            _set_text(doc, FIELD_AI_LAST_ERROR, "Document non CV")
            _set_statut(doc, status_rejected or "Rejecté")
            doc.custom_justification = f"Rejeté: ce document n'est pas un CV ({cv_check.get('reason','')})."
            doc.custom_matching_score = 0
            doc.applicant_rating = 0.0
            _save_and_reload(doc)
            return

        # --- GEMINI EXTRACTION DU CV ---
        prompt1 = """
            Rôle: Tu es un(e) recruteur(se) technique senior + analyste CV.
            Contexte: Le premier contenu fourni est le CV du candidat (PDF/image convertie en PDF). Ignore toute mise en page; analyse uniquement le texte.

            Objectif: Extraire les données CANDIDAT en un JSON strict, propre, normalisé et exploitable par un ATS. Tu dois OBLIGATOIREMENT retourner EXACTEMENT la structure ci-dessous (mêmes clés, sans rien ajouter ni retirer).

            Contraintes générales (IMPORTANTES) :
            - Réponds UNIQUEMENT par le JSON final, sans texte autour, sans balises ``` ni commentaires.
            - Ne fais AUCUNE référence au PDF, à des pages, ni à l'interface.
            - Pas d'hallucination: n'invente pas d'entreprises/diplômes non présents. Si une info est absente et non déductible, mets null ou "" (vide). Tu peux estimer UNIQUEMENT si l'indice est fort (ex: "3 ans d'expérience" mentionné explicitement).
            - Normalise l'orthographe, supprime doublons, et rends les noms propres avec capitalisation correcte.
            - Langue de sortie: FRANÇAIS (sauf noms d'outils/technos qui gardent leur orthographe canonique).
            - Respecte les types:
            - age: entier (ou null si inconnu/non déductible)
            - annee_experience: entier (approximation prudente permise si des indices explicites existent)
            - phone: liste de chaînes, format international si possible (+225…, sinon version la plus propre)
            - annee (dans expériences/diplômes): année sur 4 chiffres sous forme de chaîne (ex: "2023")
            - Distinction claire:
            - competences = compétences métiers / hard skills (ex: "gestion de projet", "data analysis")
            - outils = langages, frameworks, plateformes, logiciels, bases de données (ex: "Python", "React", "MySQL", "SAP")
            - Canonicalise les synonymes évidents: js→JavaScript, ts→TypeScript, node→Node.js, react→React.js, c sharp→C#, ms office→Microsoft Office, etc.
            - Filtrage:
            - Retire les termes trop génériques "informatique", "web", "bureautique" si non pertinents.
            - Limite competences et outils aux éléments pertinents et non redondants (max ~15 chacun), triés par pertinence (récence + fréquence + adéquation poste).
            - Expérience (experience_professionnelle):
            - Liste d'entrées {annee, titre, description} en ordre anté-chronologique (du plus récent au plus ancien).
            - "annee" = année de DÉBUT du poste si période connue (ex: 2021–2023 ⇒ "2021"), sinon l'année la plus mentionnée pour ce rôle.
            - titre = intitulé de poste normalisé (ex: "Développeur Python", "Comptable stagiaire")
            - description = 1–2 phrases synthétiques (missions/impacts/outils clés).
            - Déduplique les postes quasi-identiques.
            - Diplômes (diplomes):
            - Normalize les diplômes et institutions.
            - level ∈ { "Graduate", "Under Graduate", "Post Graduate" } avec mapping:
                - BTS/DUT/DEUG/Associate ≤ Bac+2 ⇒ "Under Graduate"
                - Licence/Bachelor/Ingénieur Bac+3/Bac+4 ⇒ "Graduate"
                - Master/MSc/MBA/Ingénieur Bac+5/Doctorat/PhD ⇒ "Post Graduate"
            - annee = année d'obtention si trouvable, sinon l'année la plus probable citée.
            - Niveau d'étude (niveau_etude): format "BAC+N" si déductible, sinon "".
            - Années d'expérience (annee_experience): calcule prudemment depuis les périodes indiquées (évite addition naïve si chevauchements); si seulement "junior/senior" est mentionné, convertis prudemment (ex: "junior"≈1–2, "senior"≈5–8), sinon 0.

            Schéma EXACT à produire (ne change pas les clés, ni la structure) :
            {
            "candidate_info": {
                "first_name": "<str ou "">",
                "last_name": "<str ou "">",
                "title": "<str ou "">",
                "age": <entier ou null>,
                "email": "<str ou "">",
                "phone": ["<str>", "..."],
                "location": "<str ou "">",
                "competences": ["<str>", "..."],
                "outils": ["<str>", "..."],
                "experience_professionnelle": [
                {
                    "annee": "<YYYY>",
                    "titre": "<str>",
                    "description": "<1-2 phrases concises: missions, résultats, outils>"
                }
                ],
                "diplomes": [
                {
                    "annee": "<YYYY ou "">",
                    "diplome": "<str>",
                    "institution": "<str ou "">",
                    "level": "Graduate" | "Under Graduate" | "Post Graduate"
                }
                ],
                "annee_experience": <entier ou null>,
                "niveau_etude": "<BAC+N ou "">"
            }
            }

            Procédure d'extraction (suivre rigoureusement) :
            1) Lire tout le texte, tolérer OCR/scan. Ignorer entêtes/pieds de page répétitifs.
            2) Identifier nom complet (first_name/last_name) même si inversé (ex: "DUPONT Marie").
            3) Extraire emails/phones/lieux même s'ils apparaissent dans l'en-tête/pied.
            4) Détecter et normaliser OUTILS vs COMPÉTENCES (voir règles ci-dessus).
            5) Construire expérience_professionnelle propre (max ~8 entrées représentatives).
            6) Diplômes: niveau + mapping "level" selon règles énoncées (très important).
            7) Déduire annee_experience si faisable, sinon null.
            8) Valider la cohérence (types, formats, années 19xx/20xx plausibles).
            9) Sortie: JSON valide, aucune clé manquante, aucune clé additionnelle.

            Rappels finaux:
            - Donne la meilleure estimation PRUDENTE quand des indices explicites existent; sinon mets null/"".
            - Respecte strictement le schéma et les énumérations.
            - Sors UNIQUEMENT le JSON final (pas de markdown).
            """

        try:
            response1 = call_gemini_with_retry(client, parts_cv + [prompt1])
        except Exception as e:
            _safe_log_error("[GEMINI] Extraction CV échouée (overloaded/404 ?)", e)
            _set_flag(doc, FLAG_MATCHING_FAILED, 1)
            _set_text(doc, FIELD_AI_LAST_ERROR, str(e))
            _set_flag(doc, FLAG_MATCHING_IN_PROGRESS, 0)
            _set_statut(doc, gemini_error_status)
            doc.custom_justification = "Analyse automatique momentanément indisponible. Traitement manuel."
            doc.custom_matching_score = 0
            doc.applicant_rating = 0.0
            _save_and_reload(doc)  # ← reload avant return
            return

        try:
            candidate_json = json.loads(response1.text)
        except Exception:
            try:
                candidate_json = json.loads(response1.text.strip("```json\n").strip("```").strip())
            except Exception as parse_err:
                raise ValueError(f"Réponse Gemini (extraction CV) non parseable: {parse_err}")

        frappe.logger().debug(f"[MATCHING] CV structuré : {json.dumps(candidate_json, ensure_ascii=False)}")

        # --- MISE À JOUR DU CANDIDAT ---
        info = candidate_json.get("candidate_info") if isinstance(candidate_json, dict) else {}
        if not isinstance(info, dict):
            info = {}

        age = info.get("age")
        if age is not None:
            doc.custom_old = int(age)

        first_name = (info.get("first_name") or "").strip()
        if first_name:
            doc.custom_first_name = first_name

        last_name = (info.get("last_name") or "").strip()
        if last_name:
            doc.custom_last_name = last_name

        annee_exp = info.get("annee_experience")
        if annee_exp is not None:
            doc.custom_minimum_experience = int(annee_exp)

        niveau = (info.get("niveau_etude") or "").strip()
        if niveau:
            doc.custom_study_level = niveau

        if hasattr(doc, "custom_outils"):
            doc.set("custom_outils", [])
            for outil in info.get("outils", []):
                doc.append("custom_outils", {"outil_name": outil})

        if hasattr(doc, "custom_skills"):
            doc.set("custom_skills", [])
            for skill in info.get("competences", []):
                doc.append("custom_skills", {"skill_name": skill})

        if hasattr(doc, "custom_assessments"):
            # Préserver les lignes déjà complétées (scores Testlify) lors d'un rematching
            completed_rows = {
                row.assessment_id: row
                for row in (doc.custom_assessments or [])
                if getattr(row, "completed", False) and getattr(row, "assessment_id", None)
            }
            doc.set("custom_assessments", [])
            for assessment in (fiche.custom_assessments or []):
                aid   = getattr(assessment, "id", None) or getattr(assessment, "assessment_id", None)
                aname = getattr(assessment, "assessment_name", None)
                if not aid:
                    continue
                if aid in completed_rows:
                    prev = completed_rows[aid]
                    doc.append("custom_assessments", {
                        "assessment_id":    aid,
                        "assessment_name":  aname or getattr(prev, "assessment_name", ""),
                        "completed":        True,
                        "assessment_score": getattr(prev, "assessment_score", 0),
                        "sent":             getattr(prev, "sent", 0),
                    })
                else:
                    doc.append("custom_assessments", {
                        "assessment_id":   aid,
                        "assessment_name": aname,
                    })

        if hasattr(doc, "custom_experiences"):
            doc.set("custom_experiences", [])
            last_annee = ""
            last_title = ""
            last_description = ""
            for experience in info.get("experience_professionnelle", []):
                annee = experience.get("annee", last_annee)
                title = experience.get("titre", last_title)
                description = experience.get("description", last_description)

                if experience.get("annee"):
                    last_annee = experience["annee"]
                if experience.get("titre"):
                    last_title = experience["titre"]
                if experience.get("description"):
                    last_description = experience["description"]

                doc.append("custom_experiences", {
                    "annee": annee,
                    "title": title,
                    "description": description
                })

        # Gestion des diplômes
        if hasattr(doc, "custom_diplomes"):
            doc.set("custom_diplomes", [])
            last_annee = ""
            last_qualification = ""
            last_institution = ""
            last_level = ""
            for diplome in info.get("diplomes", []):
                annee = diplome.get("annee", last_annee)
                qualification = diplome.get("diplome", last_qualification)
                institution = diplome.get("institution", last_institution)
                level = diplome.get("level", last_level)

                if diplome.get("annee"):
                    last_annee = diplome["annee"]
                if diplome.get("diplome"):
                    last_qualification = diplome["diplome"]
                if diplome.get("institution"):
                    last_institution = diplome["institution"]
                if diplome.get("level"):
                    last_level = diplome["level"]

                doc.append("custom_diplomes", {
                    "annee": annee,
                    "qualification": qualification,
                    "institution": institution,
                    "level": level
                })

        _save_and_reload(doc)  # ← reload après save des données CV extraites

        job_json = {
            "skills":             [getattr(r, "skill", "") for r in (fiche.custom_skills or [])],
            "outils":             [getattr(r, "outil", "") for r in (fiche.custom_outils or [])],
            "minimum_experience": fiche.custom_minimum_experience,
            "study_level":        fiche.custom_study_level,
            "fiche":              fiche.description,
        }

        prompt2 = f"""
            Rôle: Tu es un(e) recruteur(se) technique senior chargé(e) d'évaluer l'adéquation CV ↔ fiche de poste de façon rigoureuse, reproductible et sans hallucination.

            Entrées:
            - Profil candidat (JSON structuré du CV)
            - Fiche de poste (JSON: skills, outils, minimum_experience, study_level, fiche [description libre])

            Objectif:
            Calculer un score sur 100 + une justification brève (1 à 5 phrases) expliquant objectivement les principaux atouts et écarts.

            ⚖️ Barème (total = 100) — critères optionnels:
            - Compétences (skills): 40 pts
            - Outils/Technologies (outils): 25 pts
            - Niveau d'études (study_level): 15 pts
            - Expérience (minimum_experience): 20 pts

            Si un critère est absent/non renseigné dans la fiche de poste, redistribue proportionnellement son poids sur les critères restants (ex.: si seuls skills et outils présents, ils pèsent 40/(40+25)=61.54% et 25/(40+25)=38.46%, puis normalisés à 100).

            🔎 Règles d'évaluation par critère (précises):
            1) Compétences (skills)
            - Normalise (minuscules, pluriels simples, accents, variantes). Déduplique.
            - Correspondances acceptées: synonymes proches (ex.: "gestion de projet" ~ "project management").
            - Priorise explicitement les **must-have** s'ils sont identifiables dans la description de la fiche (mots-clés: "obligatoire", "indispensable", "requis", "must-have").
            - Bonus SI la compétence est **démontrée** dans des projets/missions proches des activités du poste (preuve par description d'expérience).
            - Score = couverture pondérée des compétences requises (plus fort poids pour must-have), avec crédit partiel pour équivalents proches.

            2) Outils/Technos (outils)
            - Équivalences acceptées: JS=JavaScript, TS=TypeScript, Node=Node.js, React=React.js, Express=Express.js, SQL~PostgreSQL/MySQL/SQL Server (selon contexte), MS Office~Microsoft Office, etc. Versions voisines acceptées si l'écosystème est identique.
            - Compte les familles/outils équivalents, mais évite le double comptage.
            - Score = couverture pondérée des outils requis + pertinence démontrée en projet.

            3) Niveau d'études (study_level)
            - Mappe les équivalences (Licence=Bachelor, Master=MS/MSc, Bac+5=M2/Ingénieur, etc.).
            - Si le candidat est en dessous du niveau requis → pénalité proportionnelle (forte si écart net).
            - Si au-dessus ou équivalent → validation simple (pas de sur-bonus).

            4) Expérience (minimum_experience)
            - Compare **années pertinentes** (même domaine/tech stack/responsabilités) au minimum requis.
            - Si < minimum: pénalité proportionnelle à l'écart.
            - Si > minimum: pas de bonus automatique sans pertinence claire (projets/secteur proches).
            - Privilégie la **récence** et la **pertinence** des missions par rapport aux activités principales du poste.

            🧭 Contexte & pertinence:
            - Utilise la description de la fiche (missions/activités) pour juger la similarité des projets vécus par le candidat (secteur, responsabilités, impact, environnement technique).
            - Aucune source externe. Toute information manquante dans CV/fiche = non satisfaite (pas d'invention).

            🧹 Normalisation/qualité:
            - Traite tout en minuscules pour matcher; garde les noms propres/technos dans leur forme canonique lors de la rédaction de la justification.
            - Évite de pénaliser deux fois le même écart.
            - Rends un score **entier** 0–100 (arrondi à l'unité).

            🧾 Sortie STRICTE (aucun texte autour, pas de markdown):
            - "score": entier [0..100]
            - "justification": 1 à 5 phrases max, en français, mentionnant:
            - 1–2 forces principales (ex.: compétences/outils alignés, projet très proche)
            - 1–2 écarts majeurs (ex.: must-have manquant, années d'expérience insuffisantes, niveau d'études inférieur)

            Données à évaluer:

            Profil candidat :
            {json.dumps(candidate_json, ensure_ascii=False)}

            Fiche de poste :
            {json.dumps(job_json, ensure_ascii=False)}

            Rends UNIQUEMENT ce JSON :
            {{
            "score": <entier entre 0 et 100>,
            "justification": "<jusqu'à 5 phrases expliquant objectivement les points forts et les écarts, sans détails superflus>"
            }}
            """

        try:
            response2 = call_gemini_with_retry(client, [prompt2])
        except Exception as e:
            _safe_log_error("[GEMINI] Matching échoué (overloaded/404 ?)", e)
            _set_flag(doc, FLAG_MATCHING_FAILED, 1)
            _set_text(doc, FIELD_AI_LAST_ERROR, str(e))
            _set_flag(doc, FLAG_MATCHING_IN_PROGRESS, 0)
            _set_statut(doc, gemini_error_status)
            doc.custom_justification = "Matching indisponible (surcharge moteur). Reprise auto ou traitement manuel."
            doc.custom_matching_score = 0
            doc.applicant_rating = 0.0
            _save_and_reload(doc)  # ← reload avant return
            return

        try:
            matching_score = json.loads(response2.text)
        except Exception:
            try:
                matching_score = json.loads(response2.text.strip("```json\n").strip("```").strip())
            except Exception as parse_err:
                raise ValueError(f"Réponse Gemini (score matching) non parseable: {parse_err}")

        score = 0  # valeur par défaut si matching_score n'est pas un dict valide
        if isinstance(matching_score, dict):
            score = matching_score.get("score", 0)
            doc.custom_matching_score = score
            doc.custom_justification = matching_score.get("justification")

            rating = max(0.0, min(1.0, score / 100))
            doc.applicant_rating = float(f"{rating:.2f}")

            if score <= rejected_score:
                _set_statut(doc, status_rejected)
            else:
                _set_statut(doc, qualified_status if score >= qualification_score_threshold else status_not_qualified)

            _save_and_reload(doc)  # ← reload après save du score de matching

        if score >= qualification_score_threshold:
            send_candidate_invite(doc, fiche.custom_assessments)
        else:
            send_candidate_not_matching_email(doc)

        frappe.logger().info(
            f"[MATCHING] Score : {score} | Statut : {doc.custom_status} | Candidat : {applicant_name}"
        )
        _matching_succeeded = True

    except Exception as e:
        _safe_log_error("[MATCHING] Erreur globale", e)
        _matching_error = str(e)

    finally:
        # Recharger d'abord pour éviter TimestampMismatchError,
        # puis appliquer les flags APRÈS le reload pour qu'ils soient bien persistés.
        try:
            doc.reload()
        except Exception:
            pass
        _set_flag(doc, FLAG_MATCHING_IN_PROGRESS, 0)
        if _matching_error:
            _set_flag(doc, FLAG_MATCHING_FAILED, 1)
            _set_text(doc, FIELD_AI_LAST_ERROR, _matching_error)
        elif _matching_succeeded:
            _set_flag(doc, FLAG_MATCHING_FAILED, 0)
            _set_text(doc, FIELD_AI_LAST_ERROR, "")
        doc.flags.ignore_permissions = True
        doc.validate_workflow = lambda: None
        doc.save(ignore_permissions=True)
        frappe.db.commit()